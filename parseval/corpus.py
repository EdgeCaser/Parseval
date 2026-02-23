"""File parsing and paragraph/sentence extraction.

Supports: .docx, .md, .txt, .text
"""

import os
import re

# Optional dependencies — handled gracefully if unavailable
try:
    import docx  # python-docx
    _DOCX_AVAILABLE = True
except ImportError:
    _DOCX_AVAILABLE = False

try:
    import chardet
    _CHARDET_AVAILABLE = True
except ImportError:
    _CHARDET_AVAILABLE = False

ALLOWED_EXTENSIONS = {".docx", ".txt", ".md", ".text"}

MIN_PARAGRAPH_CHARS = 20  # Shorter than this is treated as a fragment


def allowed_file(filename: str) -> bool:
    ext = os.path.splitext(filename.lower())[1]
    return ext in ALLOWED_EXTENSIONS


def extract_text(filepath: str, filename: str) -> str:
    """Extract plain text from a file. Returns the full text as a string."""
    ext = os.path.splitext(filename.lower())[1]
    if ext == ".docx":
        return _extract_docx(filepath)
    elif ext == ".md":
        return _extract_markdown(filepath)
    elif ext in (".txt", ".text"):
        return _extract_plaintext(filepath)
    else:
        raise ValueError(f"Unsupported file type: {ext}")


def split_paragraphs(text: str, min_chars: int = MIN_PARAGRAPH_CHARS) -> list:
    """Split text into paragraphs, filtering out very short ones."""
    # Normalize line endings
    text = text.replace("\r\n", "\n").replace("\r", "\n")
    # Split on one or more blank lines
    raw_paragraphs = re.split(r"\n{2,}", text)
    paragraphs = []
    for para in raw_paragraphs:
        para = para.strip()
        if len(para) >= min_chars:
            paragraphs.append(para)
    return paragraphs


def split_paragraphs_all(text: str) -> list:
    """Split text into paragraphs without filtering short ones (for analysis docs)."""
    text = text.replace("\r\n", "\n").replace("\r", "\n")
    raw_paragraphs = re.split(r"\n{2,}", text)
    return [p.strip() for p in raw_paragraphs if p.strip()]


def split_sentences(text: str) -> list:
    """Split text into sentences using spaCy if available, else regex fallback."""
    nlp = _get_nlp()
    if nlp:
        doc = nlp(text)
        return [sent.text.strip() for sent in doc.sents if sent.text.strip()]
    else:
        # Regex fallback: split on sentence-ending punctuation followed by whitespace
        sentences = re.split(r"(?<=[.!?])\s+", text)
        return [s.strip() for s in sentences if s.strip()]


# --- Private file parsers ---

def _extract_docx(filepath: str) -> str:
    if not _DOCX_AVAILABLE:
        raise ImportError("python-docx is required to read .docx files. Run: pip install python-docx")
    doc = docx.Document(filepath)
    parts = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            parts.append(text)
    # Also extract table cells
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                text = cell.text.strip()
                if text and text not in parts:
                    parts.append(text)
    return "\n\n".join(parts)


def _extract_markdown(filepath: str) -> str:
    with open(filepath, "r", encoding="utf-8", errors="replace") as f:
        text = f.read()
    # Remove fenced code blocks
    text = re.sub(r"```[\s\S]*?```", "", text)
    text = re.sub(r"`[^`\n]+`", "", text)
    # Remove headers
    text = re.sub(r"^#{1,6}\s+", "", text, flags=re.MULTILINE)
    # Remove bold/italic
    text = re.sub(r"\*{1,2}([^*\n]+)\*{1,2}", r"\1", text)
    text = re.sub(r"_{1,2}([^_\n]+)_{1,2}", r"\1", text)
    # Remove images
    text = re.sub(r"!\[.*?\]\(.*?\)", "", text)
    # Remove links (keep link text)
    text = re.sub(r"\[([^\]]+)\]\(.*?\)", r"\1", text)
    # Remove horizontal rules
    text = re.sub(r"^[-*_]{3,}\s*$", "", text, flags=re.MULTILINE)
    # Remove blockquote markers
    text = re.sub(r"^>\s*", "", text, flags=re.MULTILINE)
    # Remove list markers
    text = re.sub(r"^[\*\-\+]\s+", "", text, flags=re.MULTILINE)
    text = re.sub(r"^\d+\.\s+", "", text, flags=re.MULTILINE)
    return text


def _extract_plaintext(filepath: str) -> str:
    with open(filepath, "rb") as f:
        raw = f.read()
    if _CHARDET_AVAILABLE:
        detected = chardet.detect(raw)
        encoding = detected.get("encoding") or "latin-1"
    else:
        encoding = "utf-8"
    return raw.decode(encoding, errors="replace")


# --- spaCy singleton ---

_nlp = None
_nlp_initialized = False


def _get_nlp():
    """Return spaCy NLP object or False if unavailable."""
    global _nlp, _nlp_initialized
    if _nlp_initialized:
        return _nlp
    _nlp_initialized = True
    try:
        import spacy
        nlp = spacy.load("en_core_web_sm", disable=["ner", "parser"])
        # Enable the sentence segmenter (disabled when parser is off)
        if "senter" not in nlp.pipe_names:
            nlp.enable_pipe("senter")
        _nlp = nlp
    except Exception:
        _nlp = False
        print(
            "[parseval] WARNING: spaCy model 'en_core_web_sm' not found or failed to load.\n"
            "  Sentence splitting will use a regex fallback.\n"
            "  Run: python -m spacy download en_core_web_sm"
        )
    return _nlp


def get_spacy_nlp():
    """Public access to the spaCy NLP object (for features.py)."""
    return _get_nlp()
